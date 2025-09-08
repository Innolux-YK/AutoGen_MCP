"""
工具函數模組
包含文檔處理、文字處理等工具函數
"""

import os
import re
import json
import shutil
import tempfile
from typing import List, Dict, Any
from pathlib import Path
from docx import Document
from PIL import Image
import jieba

def extract_text_from_document(file_path: str) -> Dict[str, Any]:
    """
    從文檔中提取文字內容
    支援 DOCX 和 DOC 格式
    
    Args:
        file_path: 文檔路徑
        
    Returns:
        包含標題、關鍵字、完整文字的字典
    """
    try:
        # 轉換 DOC 為 DOCX（如果需要）
        if file_path.lower().endswith('.doc'):
            file_path = auto_convert_doc_to_docx(file_path)
            if not file_path:
                return {"title": "", "keywords": [], "full_text": ""}
        
        doc = Document(file_path)
        
        # 提取標題（第一個段落通常是標題）
        title = ""
        if doc.paragraphs:
            title = doc.paragraphs[0].text.strip()
        
        # 提取全部文字
        full_text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text += paragraph.text + "\n"
        
        # 提取關鍵字（使用jieba分詞）
        keywords = extract_keywords(full_text)
        
        return {
            "title": title,
            "keywords": keywords,
            "full_text": full_text.strip()
        }
        
    except Exception as e:
        print(f"提取文檔內容時發生錯誤: {e}")
        return {"title": "", "keywords": [], "full_text": ""}

def extract_images_from_document(file_path: str, output_dir: str) -> List[str]:
    """
    從文檔中提取圖片
    
    Args:
        file_path: 文檔路徑
        output_dir: 圖片輸出目錄
        
    Returns:
        提取的圖片路徑列表
    """
    try:
        # 確保輸出目錄存在
        os.makedirs(output_dir, exist_ok=True)
        
        # 轉換 DOC 為 DOCX（如果需要）
        if file_path.lower().endswith('.doc'):
            file_path = auto_convert_doc_to_docx(file_path)
            if not file_path:
                return []
        
        doc = Document(file_path)
        image_paths = []
        
        # 從文檔中提取圖片
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                img_data = rel.target_part.blob
                
                # 生成圖片檔名
                filename = os.path.basename(file_path).replace('.docx', '')
                img_filename = f"{filename}_img_{len(image_paths)+1}.png"
                img_path = os.path.join(output_dir, img_filename)
                
                # 保存圖片
                with open(img_path, 'wb') as f:
                    f.write(img_data)
                
                image_paths.append(img_path)
        
        return image_paths
        
    except Exception as e:
        print(f"提取圖片時發生錯誤: {e}")
        return []

def auto_convert_doc_to_docx(doc_path: str, permanent: bool = False) -> str:
    """
    將 DOC 檔案轉換為 DOCX
    
    Args:
        doc_path: DOC 檔案路徑
        permanent: 是否永久轉換（覆蓋原檔案）
        
    Returns:
        轉換後的 DOCX 檔案路徑
    """
    try:
        # 這裡需要實現 DOC 到 DOCX 的轉換
        # 由於 python-docx 不直接支援 DOC，這裡返回原路徑
        # 實際應用中可能需要使用 python-docx2 或其他工具
        docx_path = doc_path.replace('.doc', '.docx')
        
        if permanent and not os.path.exists(docx_path):
            # 這裡應該實現真正的轉換邏輯
            # 暫時複製檔案
            shutil.copy2(doc_path, docx_path)
        
        return docx_path if os.path.exists(docx_path) else doc_path
        
    except Exception as e:
        print(f"轉換 DOC 檔案時發生錯誤: {e}")
        return doc_path

def chunk_text(text: str, chunk_size: int = 500, chunk_overlap: int = 50) -> List[str]:
    """
    將文字分割成塊
    
    Args:
        text: 要分割的文字
        chunk_size: 塊大小
        chunk_overlap: 重疊大小
        
    Returns:
        文字塊列表
    """
    if not text:
        return []
    
    # 按句子分割
    sentences = re.split(r'[。！？\n]', text)
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
            
        # 如果當前塊加上新句子超過大小限制
        if len(current_chunk) + len(sentence) > chunk_size:
            if current_chunk:
                chunks.append(current_chunk.strip())
                # 保留重疊部分
                if chunk_overlap > 0:
                    current_chunk = current_chunk[-chunk_overlap:]
                else:
                    current_chunk = ""
        
        current_chunk += sentence + "。"
    
    # 添加最後一塊
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    return chunks

def clean_text(text: str) -> str:
    """
    清理文字內容
    
    Args:
        text: 要清理的文字
        
    Returns:
        清理後的文字
    """
    if not text:
        return ""
    
    # 移除多餘的空白字元
    text = re.sub(r'\s+', ' ', text)
    
    # 移除特殊字元（保留中文、英文、數字、常用標點）
    text = re.sub(r'[^\u4e00-\u9fff\w\s，。！？；：「」『』（）\[\]{}.,;:!?()\-\'"]+', '', text)
    
    return text.strip()

def extract_keywords(text: str, max_keywords: int = 10) -> List[str]:
    """
    提取關鍵字
    
    Args:
        text: 文字內容
        max_keywords: 最大關鍵字數量
        
    Returns:
        關鍵字列表
    """
    if not text:
        return []
    
    # 使用jieba分詞
    words = jieba.cut(text)
    
    # 過濾停用詞和短詞
    stopwords = {'的', '是', '在', '了', '有', '和', '與', '或', '但是', '然而', '因為', '所以'}
    keywords = []
    
    for word in words:
        word = word.strip()
        if (len(word) > 1 and 
            word not in stopwords and 
            not word.isdigit() and
            not re.match(r'^[a-zA-Z]+$', word)):
            keywords.append(word)
    
    # 去重並限制數量
    keywords = list(set(keywords))[:max_keywords]
    
    return keywords

def save_metadata(metadata: List[Dict], filepath: str):
    """
    保存元數據到JSON檔案
    
    Args:
        metadata: 元數據列表
        filepath: 檔案路徑
    """
    try:
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"保存元數據時發生錯誤: {e}")

def cleanup_temp_files(directory: str):
    """
    清理臨時檔案
    
    Args:
        directory: 要清理的目錄
    """
    try:
        for root, dirs, files in os.walk(directory):
            for file in files:
                if file.startswith('~$') or file.endswith('.tmp'):
                    temp_file = os.path.join(root, file)
                    try:
                        os.remove(temp_file)
                        print(f"已清理臨時檔案: {file}")
                    except:
                        pass
    except Exception as e:
        print(f"清理臨時檔案時發生錯誤: {e}")
